from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from typing import ClassVar

import pytest
from docx import Document as DocxDocument

from personal_rag.config import Settings
from personal_rag.errors import RagError
from personal_rag.parsers import DocumentParser, safe_display_name


@pytest.fixture
def settings(tmp_path: Path) -> Settings:
    return Settings(
        auth_enabled=False,
        data_dir=tmp_path / "data",
        upload_max_bytes=64 * 1024,
        max_pdf_pages=3,
        max_extracted_characters=1000,
    )


def assert_error(error: pytest.ExceptionInfo[RagError], code: str) -> None:
    assert error.value.code == code


def test_safe_display_name_removes_paths_controls_and_bounds_length() -> None:
    assert safe_display_name("../../private/notes.md") == "notes.md"
    assert safe_display_name(r"C:\fakepath\notes.txt") == "notes.txt"
    assert safe_display_name("\x00\n report .md ") == "report .md"

    long_name = safe_display_name(f"{'a' * 300}.pdf")
    assert len(long_name) <= 255
    assert long_name.endswith(".pdf")


def test_markdown_is_split_into_scalar_metadata_sections(
    tmp_path: Path, settings: Settings
) -> None:
    path = tmp_path / "notes.md"
    path.write_text(
        "# Alpha\nFirst section.\n\n```md\n# Not a heading\n```\n\n## Beta\nSecond section.",
        encoding="utf-8",
    )

    documents = DocumentParser(settings).parse(path, display_name="../My Notes.md")

    assert [document.metadata["section"] for document in documents] == ["Alpha", "Beta"]
    assert "Not a heading" in documents[0].text
    assert documents[0].metadata["source_name"] == "My Notes.md"
    assert all(
        isinstance(value, str | int | float | bool)
        for document in documents
        for value in document.metadata.values()
    )


def test_text_parser_rejects_binary_empty_and_character_limit(
    tmp_path: Path, settings: Settings
) -> None:
    parser = DocumentParser(settings)
    binary = tmp_path / "binary.txt"
    binary.write_bytes(b"hello\x00world")
    with pytest.raises(RagError) as binary_error:
        parser.parse(binary)
    assert_error(binary_error, "invalid_file_signature")

    empty = tmp_path / "empty.txt"
    empty.write_text(" \n\t", encoding="utf-8")
    with pytest.raises(RagError) as empty_error:
        parser.parse(empty)
    assert_error(empty_error, "empty_document")

    too_long = tmp_path / "long.txt"
    too_long.write_text("x" * 1001, encoding="utf-8")
    with pytest.raises(RagError) as long_error:
        parser.parse(too_long)
    assert_error(long_error, "character_limit_exceeded")


def test_parser_rejects_unsupported_spoofed_oversized_and_symlinked_files(
    tmp_path: Path, settings: Settings
) -> None:
    parser = DocumentParser(settings)
    unsupported = tmp_path / "notes.csv"
    unsupported.write_text("a,b", encoding="utf-8")
    with pytest.raises(RagError) as unsupported_error:
        parser.parse(unsupported)
    assert_error(unsupported_error, "unsupported_file_type")

    spoofed = tmp_path / "spoofed.pdf"
    spoofed.write_text("not a pdf", encoding="utf-8")
    with pytest.raises(RagError) as spoofed_error:
        parser.parse(spoofed)
    assert_error(spoofed_error, "invalid_file_signature")

    oversized = tmp_path / "oversized.txt"
    oversized.write_bytes(b"x" * (settings.upload_max_bytes + 1))
    with pytest.raises(RagError) as oversized_error:
        parser.parse(oversized)
    assert_error(oversized_error, "file_too_large")

    target = tmp_path / "target.txt"
    target.write_text("hello", encoding="utf-8")
    symlink = tmp_path / "linked.txt"
    symlink.symlink_to(target)
    with pytest.raises(RagError) as symlink_error:
        parser.parse(symlink)
    assert_error(symlink_error, "unsafe_file")


def test_docx_is_split_by_heading_and_includes_tables(tmp_path: Path, settings: Settings) -> None:
    path = tmp_path / "report.docx"
    source = DocxDocument()
    source.add_heading("Overview", level=1)
    source.add_paragraph("A short summary.")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    source.add_heading("Details", level=2)
    source.add_paragraph("More detail.")
    source.save(path)

    documents = DocumentParser(settings).parse(path)

    assert [document.metadata["section"] for document in documents] == [
        "Overview",
        "Details",
    ]
    assert "Key | Value" in documents[0].text
    assert "More detail" in documents[1].text


def test_encrypted_office_container_is_reported_explicitly(
    tmp_path: Path, settings: Settings
) -> None:
    path = tmp_path / "protected.docx"
    path.write_bytes(bytes.fromhex("D0CF11E0A1B11AE1") + b"encrypted")

    with pytest.raises(RagError) as error:
        DocumentParser(settings).parse(path)

    assert_error(error, "encrypted_document")


def test_pdf_emits_one_document_per_text_page(
    tmp_path: Path, settings: Settings, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = tmp_path / "paper.pdf"
    path.write_bytes(b"%PDF-1.7\n")

    class FakeReader:
        is_encrypted = False
        pages: ClassVar[list[SimpleNamespace]] = [
            SimpleNamespace(extract_text=lambda: "Page one", images=[]),
            SimpleNamespace(extract_text=lambda: "  ", images=[]),
            SimpleNamespace(extract_text=lambda: "Page three", images=[]),
        ]

        def __init__(self, _path: Path, *, strict: bool = False) -> None:
            assert strict is False

    monkeypatch.setattr("personal_rag.parsers.PdfReader", FakeReader)
    documents = DocumentParser(settings).parse(path)

    assert [document.metadata["page_number"] for document in documents] == [1, 3]
    assert [document.text for document in documents] == ["Page one", "Page three"]


@pytest.mark.parametrize(
    ("reader", "code"),
    [
        (SimpleNamespace(is_encrypted=True, pages=[]), "encrypted_document"),
        (
            SimpleNamespace(
                is_encrypted=False,
                pages=[SimpleNamespace(extract_text=lambda: "", images=[object()])],
            ),
            "image_only_document",
        ),
        (
            SimpleNamespace(
                is_encrypted=False,
                pages=[SimpleNamespace(extract_text=lambda: "", images=[])],
            ),
            "empty_document",
        ),
        (
            SimpleNamespace(
                is_encrypted=False,
                pages=[SimpleNamespace(extract_text=lambda: "x", images=[])] * 4,
            ),
            "page_limit_exceeded",
        ),
    ],
)
def test_pdf_reports_encrypted_image_only_empty_and_page_limit(
    tmp_path: Path,
    settings: Settings,
    monkeypatch: pytest.MonkeyPatch,
    reader: SimpleNamespace,
    code: str,
) -> None:
    path = tmp_path / "paper.pdf"
    path.write_bytes(b"%PDF-1.7\n")
    monkeypatch.setattr("personal_rag.parsers.PdfReader", lambda *_args, **_kwargs: reader)

    with pytest.raises(RagError) as error:
        DocumentParser(settings).parse(path)

    assert_error(error, code)
